import os
import io
from PIL import Image
from pypdf import PdfWriter, PdfReader
import fitz  # PyMuPDF

import pytesseract
from docx import Document

def merge_pdfs(file_paths):
    """
    Merges multiple PDFs into a single temporary PDF stream (BytesIO) or file.
    Returns a BytesIO object containing the merged PDF.
    """
    merger = PdfWriter()
    for path in file_paths:
        merger.append(path)
    
    output_stream = io.BytesIO()
    merger.write(output_stream)
    output_stream.seek(0)
    return output_stream

def convert_to_images(pdf_stream_or_path):
    """
    Converts a PDF (given as path or stream) to a list of PIL Images using PyMuPDF (fitz).
    This avoids the reliance on external 'poppler' tools.
    """
    images = []
    
    # Open the PDF
    if isinstance(pdf_stream_or_path, io.BytesIO):
        # PyMuPDF expects bytes or filename. stream should be bytes.
        # fitz.open(stream=..., filetype="pdf")
        pdf_stream_or_path.seek(0)
        doc = fitz.open(stream=pdf_stream_or_path.read(), filetype="pdf")
    else:
        doc = fitz.open(pdf_stream_or_path)

    for page_index in range(len(doc)):
        page = doc[page_index]
        # Render page to an image (pixmap)
        # matrix=fitz.Matrix(2, 2) makes it higher resolution (2x zoom) for better OCR
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        
        # Convert to PIL Image
        # Pixmap.samples is the raw byte data
        mode = "RGBA" if pix.alpha else "RGB"
        img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
        
        # Ensure RGB for OCR/Saving
        if mode == "RGBA":
            img = img.convert("RGB")
            
        images.append(img)
        
    return images

def perform_ocr(images, progress_callback=None):
    """
    Performs OCR on a list of PIL Images using pytesseract (Tesseract OCR).
    Returns a list of strings (one per page).
    
    progress_callback: function(current, total)
    """
    extracted_text = []
    total_pages = len(images)
    
    for i, img in enumerate(images):
        try:
            # Tesseract works directly with PIL images
            text = pytesseract.image_to_string(img)
            extracted_text.append(text.strip())
            
        except Exception as e:
            error_msg = f"[OCR Error on page {i+1}: {e}]"
            print(error_msg)
            # Check for common "tesseract not installed" error
            if "tesseract is not installed" in str(e).lower() or "not found" in str(e).lower():
                extracted_text.append("Error: Tesseract is not installed on the system. Please install it (apt-get install tesseract-ocr).")
            else:
                extracted_text.append(error_msg)
        
        if progress_callback:
            progress_callback(i + 1, total_pages)
            
    return extracted_text

def save_to_docx(text_data, output_path):
    """
    Saves the list of text pages to a .docx file.
    """
    doc = Document()
    doc.add_heading('Extracted Text', 0)

    for i, page_text in enumerate(text_data):
        doc.add_heading(f'Page {i+1}', level=1)
        doc.add_paragraph(page_text)
        doc.add_page_break()

    doc.save(output_path)
